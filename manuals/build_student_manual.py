"""Build the Student User Manual DOCX from markdown chapter files.

Usage:
    python manuals/build_student_manual.py

Generates: manuals/student-user-manual.docx
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

MANUAL_DIR = Path(__file__).parent
STUDENT_DIR = MANUAL_DIR / "student"
OUTPUT = MANUAL_DIR / "student-user-manual.docx"

# Chapter files in order
CHAPTERS = [
    "ch01_getting_started.md",
    "ch02_navigating_docs.md",
    "ch03_first_measurement.md",
    "ch04_running_examples.md",
    "ch05_lab_workflow.md",
    "ch06_variables_calc_logging.md",
    "ch07_scripting.md",
    "ch08_multi_instrument.md",
    "ch09_data_analysis.md",
    "ch10_tips_gotchas.md",
    "ch11_troubleshooting.md",
]


def _set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, text):
    """Add a code block with gray background and monospace font."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # Gray background
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F0F0F0")
        shading.set(qn("w:val"), "clear")
        p.paragraph_format.element.get_or_add_pPr().append(shading)


def add_table(doc, headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()  # spacing after table


def parse_markdown(doc, md_text):
    """Parse markdown text and add to the DOCX document."""
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Indented code blocks (4 spaces)
        if line.startswith("    ") and not line.strip().startswith("-") and not line.strip().startswith("|"):
            code_block = []
            while i < len(lines) and lines[i].startswith("    "):
                code_block.append(lines[i][4:])
                i += 1
            add_code_block(doc, "\n".join(code_block))
            continue

        # Tables
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True
                table_headers = [c.strip() for c in line.strip().strip("|").split("|")]
                i += 1  # skip separator line
                if i < len(lines) and "---" in lines[i]:
                    i += 1
                table_rows = []
                continue
            else:
                row = [c.strip() for c in line.strip().strip("|").split("|")]
                table_rows.append(row)
                i += 1
                # Check if next line is still a table
                if i >= len(lines) or "|" not in lines[i] or not lines[i].strip().startswith("|"):
                    add_table(doc, table_headers, table_rows)
                    in_table = False
                continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Bold text paragraphs
        stripped = line.strip()
        if stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
            i += 1
            continue

        # Regular paragraph - collect consecutive non-empty lines
        para_lines = []
        while i < len(lines):
            l = lines[i]
            if not l.strip():
                break
            if l.startswith("#"):
                break
            if l.startswith("    ") and not l.strip().startswith("-"):
                break
            if l.strip().startswith("|") and "|" in l:
                break
            if l.strip().startswith("```"):
                break
            para_lines.append(l.strip())
            i += 1

        if para_lines:
            text = " ".join(para_lines)
            # Handle inline code
            p = doc.add_paragraph()
            parts = re.split(r"(`[^`]+`)", text)
            for part in parts:
                if part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                elif part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Handle bold within text
                    bold_parts = re.split(r"(\*\*[^*]+\*\*)", part)
                    for bp in bold_parts:
                        if bp.startswith("**") and bp.endswith("**"):
                            run = p.add_run(bp[2:-2])
                            run.bold = True
                        else:
                            p.add_run(bp)
        continue

    # Flush any remaining table
    if in_table and table_headers:
        add_table(doc, table_headers, table_rows)


def build():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SCPI Instrument Toolkit")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x50, 0x0, 0x0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student User Manual")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x50, 0x0, 0x0)

    doc.add_paragraph()

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("ESET 453 - Validation and Verification")
    run.font.size = Pt(14)

    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("Texas A&M University")
    run.font.size = Pt(14)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    doc.add_heading("Table of Contents", level=1)
    toc = doc.add_paragraph()
    toc.add_run(
        "Chapter 1: Getting Started\n"
        "Chapter 2: Navigating the Official Docs\n"
        "Chapter 3: Your First Measurement\n"
        "Chapter 4: Finding and Running Examples\n"
        "Chapter 5: Lab Report Workflow\n"
        "Chapter 6: Variables, Calc, and Logging\n"
        "Chapter 7: Scripting Mastery\n"
        "Chapter 8: Multi-Instrument Workflows\n"
        "Chapter 9: Data Analysis\n"
        "Chapter 10: Instrument Tips and Gotchas\n"
        "Chapter 11: Troubleshooting"
    )
    doc.add_page_break()

    # ── Chapters ──
    for ch_file in CHAPTERS:
        path = STUDENT_DIR / ch_file
        if not path.exists():
            print(f"  WARNING: {ch_file} not found, skipping")
            continue
        print(f"  Adding {ch_file}...")
        md = path.read_text()
        parse_markdown(doc, md)
        doc.add_page_break()

    # ── Save ──
    doc.save(str(OUTPUT))
    print(f"\nGenerated: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
